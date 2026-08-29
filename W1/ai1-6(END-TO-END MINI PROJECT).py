#=======================================
"""
---Why is this better than the one i made?---
First of all everything is STRUCTURED u can see 
from resume to job description analysis and everything.
Everything is separate from getting the match results and all
to scoring and giving details.

---What i did wrong---
Should make it structured like this one.
(like Resume, Job, MatchResult, etc)
Functions for doing separate things.
Functions for importing documents.
Rate of calling the llm is limited.


---Final Verdict and things to remember---
Dont stress.
Things should be structured.
Separate function for each functionality.
Nothing more nothing less.

maybe i will make a diff proj using the llm? lets see.
"""


import os
from pathlib import Path
from dotenv import load_dotenv
from pydantic import BaseModel
from groq import Groq
import json
from pypdf import PdfReader
from docx import Document

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")
model = "openai/gpt-oss-120b"
client = Groq(api_key = my_api_key)


class Job(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    educational_requirements: list[str]
    responsibilities: list[str]
job_schema = Job.model_json_schema()

job_description = """
Description
Software Development Engineer 6 months Internship (In-Person)

Introduction
At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. The intense focus we have on our customers is why we are one of the world’s most beloved brands – customer obsession is part of our company DNA. Our interns write real software and collaborate with a select group of experienced software development engineers (SDEs) who guide interns on projects that matter to our customers. As an intern, you will be matched to a manager and a mentor. You will have the opportunity to influence the evolution of Amazon technology and lead mission critical projects early in your career. Your design, code, and raw smarts will contribute to solving some of the most complex technical challenges in the areas of distributed systems, data mining, automation, optimization, scalability, and security – just to name a few. In addition to working on an impactful project, you will have the opportunity to engage with Amazonians for both personal and professional development, expand your network, and participate in activities with other interns throughout your internship. No matter the location of your internship, we give you the tools to own your project and learn in a real-world setting. Many of our technologies overlap, and you would be hard pressed to find a team that is not using Amazon Web Services (AWS), touching the catalogue, or iterating services to better personalize for customers. If this opportunity interests you, apply and come chart your own path at Amazon.

Job location
By applying to this position your application will be considered for all locations we hire for in India. This includes but is not limited to Bengaluru, Chennai, Hyderabad, Delhi, Mumbai and Pune.

Duration and Timelines
Amazon internships are full-time (40 hours/week) for 24 - 26 consecutive weeks (six months) with start dates between January - June 

Please note that Amazon internships require full-time commitment during the duration of the internship. During the course of internship, interns should not have any conflicts including but not limited to academic projects, classes or other internships/employment. Any exam related details must be shared with the hiring manager to plan for absence during those days. Specific team norms around working hours will be communicated by the hiring/ reporting manager at the time of commencement of internship. Candidates receiving internship will be required to submit declaration of their availability to complete the entire duration of internship duly signed by a competent authority at their University. Internship offer will be subjected to successful submission of the declaration.

Key job responsibilities
• Collaborate with experienced cross-disciplinary Amazonians to conceive, design, and bring innovative products and services to market.
• Design and build innovative technologies in a large distributed computing environment, and help lead fundamental changes in the industry.
• Create solutions to run predictions on distributed systems with exposure to innovative
technologies at incredible scale and speed.
• Build distributed storage, index, and query systems that are scalable, fault-tolerant, low cost, and easy to manage/use.
• Ability to design and code the right solutions starting with broadly defined problems.
• Work in an agile environment to deliver high-quality software.

Basic Qualifications
- Knowledge of computer science fundamentals such as object-oriented design, operating systems, algorithms, data structures, and complexity analysis
- Knowledge of programming languages such as C/C++, Python, Java or Perl
- Currently enrolled in a Bachelor's or Master's Degree in Computer Science, Computer Engineering, or related field at time of application -Year of Graduation 2027

Preferred Qualifications
- Previous technical internship(s).
- Experience with distributed, multi-tiered systems, algorithms, and relational databases.
- Experience in optimization mathematics such as linear programming and nonlinear optimization.
- Effectively articulate technical challenges and solutions.
- Adept at handling ambiguous or undefined problems as well as ability to think abstractly.
"""

system_message = f"""
You are a Senior HR Manager with over 50 years of experience.
Your job is to analyze the job description and extract
structure information from them.

Return only valid JSON matching this schema:
{job_schema}

IMPORTANT:
DO NOT RETURN THE SCHEMA ITSELF.
DO NOT RETURN FIELDS LIKE "Properties", "title" or "type".
FILL ALL THE SCHEMA WITH ACTUAL INFO FROM THE JOB DESCRIPTION.

If minimum experience is not mentioned return null
If information for a list is missing return an empty list
DO NOT INVENT INFORMATION.
"""

user_message = f"""
Analyze the following JD.

{job_description}
"""

system_message = {
    "role": "system",
    "content": system_message
}
message_user = {
    "role": "user",
    "content": user_message
}
response_format = {
    "type": "json_object"
}


messages = [system_message, message_user]
response = client.chat.completions.create(model= model, messages= messages, response_format= response_format)
ans = response.choices[0].message.content
raw_json = ans

job_data = json.loads(raw_json)
job = Job(**job_data)




class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str
    role: str
    duration: str
    description: str
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str
    email: str
    phone_no: str
    total_experience_in_years: str | None
    skills: list[str]
    experience: list[Experience]
    education: list[str]
    projects: list[str]
    certifications: list[str]

resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.
    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None


import time
resume_folder = Path(r"D:\AI Engineer\RESUMES")
all_results=[]
for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) 
    time.sleep(15)
    result = final_score(job, parsed_resume)
    time.sleep(15)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)

print(all_results)
# top_2 = all_results[:2]
# worst_2 = all_results[-2:]


# print("TOP 2 CANDIDATES")
# for candidate in top_2:

#     print(
#         candidate["name"],
#         "-",
#         candidate["score"],
#         "%"
#     )

#     print(candidate["details"])

# print("LOWEST 2 CANDIDATES")
# for candidate in worst_2:

#     print(
#         candidate["name"],
#         "-",
#         candidate["score"],
#         "%"
#     )
#     print(candidate["details"])